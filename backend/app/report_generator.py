"""
Professional After Action Report Generator
Generates Word (.docx) and PDF reports for cybersecurity incident response tabletop exercises.
"""
from datetime import datetime
from typing import Dict, Any
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY


def get_risk_color_hex(risk_rating: str) -> str:
    """Get hex color for risk rating"""
    colors_map = {
        "Critical": "#DC2626",  # Red
        "High": "#EA580C",      # Orange
        "Medium": "#EAB308",    # Yellow
        "Low": "#22C55E",       # Green
        "Very Low": "#3B82F6",  # Blue
    }
    return colors_map.get(risk_rating, "#6B7280")


def get_risk_color_rgb(risk_rating: str) -> tuple:
    """Get RGB color for risk rating"""
    hex_color = get_risk_color_hex(risk_rating)
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def generate_word_report(report_data: Dict[str, Any]) -> BytesIO:
    """
    Generate a professional Word document (.docx) from the AAR data.
    Returns a BytesIO object containing the document.
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('After Action Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Cybersecurity Incident Response Tabletop Exercise')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary Section
    doc.add_heading('Executive Summary', 1)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Scenario: ').bold = True
    summary_para.add_run(report_data.get('scenario_name', 'Unknown Scenario'))
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Report Generated: ').bold = True
    summary_para.add_run(datetime.fromisoformat(report_data.get('generated_at', datetime.now().isoformat())).strftime('%B %d, %Y at %I:%M %p'))
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Overall Risk Rating: ').bold = True
    risk_run = summary_para.add_run(report_data.get('overall_risk_rating', 'Not Rated'))
    risk_run.font.color.rgb = RGBColor(*get_risk_color_rgb(report_data.get('overall_risk_rating', 'Not Rated')))
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Average Effectiveness Score: ').bold = True
    summary_para.add_run(f"{report_data.get('overall_risk_score', 0):.2f}/10")
    
    doc.add_paragraph()  # Spacing
    
    # Phase-by-Phase Analysis
    doc.add_heading('Phase-by-Phase Analysis', 1)
    
    phase_analyses = report_data.get('phase_analyses', [])
    for idx, phase in enumerate(phase_analyses, 1):
        # Phase Header
        phase_heading = doc.add_heading(f"Phase {phase.get('phase_order', idx) + 1}: {phase.get('phase_name', 'Unknown Phase')}", 2)
        
        # Phase Details
        details_para = doc.add_paragraph()
        details_para.add_run('Risk Rating: ').bold = True
        risk_run = details_para.add_run(phase.get('risk_rating', 'Not Rated'))
        risk_run.font.color.rgb = RGBColor(*get_risk_color_rgb(phase.get('risk_rating', 'Not Rated')))
        
        if phase.get('average_rating'):
            details_para = doc.add_paragraph()
            details_para.add_run('Average Effectiveness Rating: ').bold = True
            details_para.add_run(f"{phase.get('average_rating'):.2f}/10")
        
        details_para = doc.add_paragraph()
        details_para.add_run('Total Responses: ').bold = True
        details_para.add_run(str(phase.get('total_responses', 0)))
        
        # GM Notes
        if phase.get('gm_notes'):
            doc.add_heading('Game Manager Notes & Takeaways', 3)
            gm_para = doc.add_paragraph(phase.get('gm_notes'))
            gm_para_format = gm_para.runs[0].font
            gm_para_format.italic = True
        
        doc.add_paragraph()  # Spacing between phases
    
    # Recommendations Section
    doc.add_heading('Recommendations', 1)
    rec_para = doc.add_paragraph(
        "Based on the phase-by-phase analysis, the following recommendations are provided "
        "to improve organizational readiness and response capabilities:"
    )
    
    # Add recommendations based on risk ratings
    high_risk_phases = [p for p in phase_analyses if p.get('risk_rating') in ['Critical', 'High']]
    if high_risk_phases:
        doc.add_heading('High Priority Recommendations', 2)
        doc.add_paragraph(
            f"Organizational effectiveness was rated as Critical or High risk in {len(high_risk_phases)} phase(s). "
            "Immediate action is recommended to address identified gaps in detection and response capabilities.",
            style='List Bullet'
        )
        doc.add_paragraph(
            "Consider implementing additional security controls, enhanced monitoring, and regular training exercises.",
            style='List Bullet'
        )
    
    medium_risk_phases = [p for p in phase_analyses if p.get('risk_rating') == 'Medium']
    if medium_risk_phases:
        doc.add_heading('Medium Priority Recommendations', 2)
        doc.add_paragraph(
            f"Organizational effectiveness was rated as Medium risk in {len(medium_risk_phases)} phase(s). "
            "Review and enhance existing security controls and response procedures.",
            style='List Bullet'
        )
    
    # Conclusion
    doc.add_heading('Conclusion', 1)
    conclusion_para = doc.add_paragraph(
        "This after action report provides a comprehensive analysis of the tabletop exercise outcomes. "
        "The risk ratings and Game Manager notes should be used to inform security improvements and training priorities. "
        "Regular tabletop exercises are recommended to maintain and improve incident response capabilities."
    )
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph('--- End of Report ---')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer_para.runs[0].font
    footer_format.size = Pt(10)
    footer_format.italic = True
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_pdf_report(report_data: Dict[str, Any]) -> BytesIO:
    """
    Generate a professional PDF document from the AAR data.
    Returns a BytesIO object containing the PDF.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1E40AF'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=14,
        textColor=colors.HexColor('#6B7280'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#1E40AF'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#374151'),
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )
    
    # Title
    elements.append(Paragraph('After Action Report', title_style))
    elements.append(Paragraph('Cybersecurity Incident Response Tabletop Exercise', subtitle_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Executive Summary
    elements.append(Paragraph('Executive Summary', heading_style))
    
    scenario_name = report_data.get('scenario_name', 'Unknown Scenario')
    generated_at = datetime.fromisoformat(report_data.get('generated_at', datetime.now().isoformat())).strftime('%B %d, %Y at %I:%M %p')
    overall_risk = report_data.get('overall_risk_rating', 'Not Rated')
    overall_score = report_data.get('overall_risk_score', 0)
    
    summary_data = [
        ['Scenario:', scenario_name],
        ['Report Generated:', generated_at],
        ['Overall Risk Rating:', overall_risk],
        ['Average Effectiveness Score:', f"{overall_score:.2f}/10"],
    ]
    
    summary_table = Table(summary_data, colWidths=[2*inch, 4*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F3F4F6')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # Phase-by-Phase Analysis
    elements.append(Paragraph('Phase-by-Phase Analysis', heading_style))
    
    phase_analyses = report_data.get('phase_analyses', [])
    for idx, phase in enumerate(phase_analyses, 1):
        phase_name = f"Phase {phase.get('phase_order', idx) + 1}: {phase.get('phase_name', 'Unknown Phase')}"
        elements.append(Paragraph(phase_name, subheading_style))
        
        # Phase details
        phase_details = []
        phase_details.append(['Risk Rating:', phase.get('risk_rating', 'Not Rated')])
        if phase.get('average_rating'):
            phase_details.append(['Average Effectiveness Rating:', f"{phase.get('average_rating'):.2f}/10"])
        phase_details.append(['Total Responses:', str(phase.get('total_responses', 0))])
        
        phase_table = Table(phase_details, colWidths=[2.5*inch, 3.5*inch])
        phase_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F9FAFB')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(phase_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # GM Notes
        if phase.get('gm_notes'):
            elements.append(Paragraph('Game Manager Notes & Takeaways', ParagraphStyle(
                'GMNotesHeading',
                parent=styles['Heading3'],
                fontSize=12,
                spaceAfter=6,
                fontName='Helvetica-Bold'
            )))
            gm_notes_para = Paragraph(f"<i>{phase.get('gm_notes')}</i>", styles['Normal'])
            elements.append(gm_notes_para)
            elements.append(Spacer(1, 0.1*inch))
        
        elements.append(Spacer(1, 0.2*inch))
    
    # Recommendations
    elements.append(PageBreak())
    elements.append(Paragraph('Recommendations', heading_style))
    
    rec_text = (
        "Based on the phase-by-phase analysis, the following recommendations are provided "
        "to improve organizational readiness and response capabilities:"
    )
    elements.append(Paragraph(rec_text, styles['Normal']))
    elements.append(Spacer(1, 0.2*inch))
    
    # High Priority Recommendations
    high_risk_phases = [p for p in phase_analyses if p.get('risk_rating') in ['Critical', 'High']]
    if high_risk_phases:
        elements.append(Paragraph('High Priority Recommendations', subheading_style))
        elements.append(Paragraph(
            f"• Organizational effectiveness was rated as Critical or High risk in {len(high_risk_phases)} phase(s). "
            "Immediate action is recommended to address identified gaps in detection and response capabilities.",
            styles['Normal']
        ))
        elements.append(Paragraph(
            "• Consider implementing additional security controls, enhanced monitoring, and regular training exercises.",
            styles['Normal']
        ))
        elements.append(Spacer(1, 0.15*inch))
    
    # Medium Priority Recommendations
    medium_risk_phases = [p for p in phase_analyses if p.get('risk_rating') == 'Medium']
    if medium_risk_phases:
        elements.append(Paragraph('Medium Priority Recommendations', subheading_style))
        elements.append(Paragraph(
            f"• Organizational effectiveness was rated as Medium risk in {len(medium_risk_phases)} phase(s). "
            "Review and enhance existing security controls and response procedures.",
            styles['Normal']
        ))
        elements.append(Spacer(1, 0.15*inch))
    
    # Conclusion
    elements.append(Paragraph('Conclusion', heading_style))
    conclusion_text = (
        "This after action report provides a comprehensive analysis of the tabletop exercise outcomes. "
        "The risk ratings and Game Manager notes should be used to inform security improvements and training priorities. "
        "Regular tabletop exercises are recommended to maintain and improve incident response capabilities."
    )
    elements.append(Paragraph(conclusion_text, styles['Normal']))
    
    # Footer
    elements.append(Spacer(1, 0.5*inch))
    footer = Paragraph('<i>--- End of Report ---</i>', ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#9CA3AF')
    ))
    elements.append(footer)
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

